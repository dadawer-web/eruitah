import asyncio
import logging
import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

logger = logging.getLogger(__name__)

HEADING_SIZES = {1: Pt(22), 2: Pt(16), 3: Pt(13), 4: Pt(12)}
HEADING_COLORS = {
    1: RGBColor(0x1A, 0x1A, 0x2E),
    2: RGBColor(0x2D, 0x2D, 0x44),
    3: RGBColor(0x44, 0x44, 0x5C),
    4: RGBColor(0x55, 0x55, 0x6B),
}
BODY_SIZE = Pt(11)
BODY_COLOR = RGBColor(0x33, 0x33, 0x33)
CODE_COLOR = RGBColor(0xC7, 0x25, 0x4E)
FONT_BODY = "Microsoft YaHei"
FONT_HEADING = "Microsoft YaHei"
FONT_CODE = "Consolas"


def _setup_document() -> Document:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    style_normal = doc.styles["Normal"]
    style_normal.font.name = FONT_BODY
    style_normal.font.size = BODY_SIZE
    style_normal.font.color.rgb = BODY_COLOR
    style_normal.paragraph_format.space_after = Pt(8)
    style_normal.paragraph_format.space_before = Pt(2)
    style_normal.paragraph_format.line_spacing = 1.5

    for level in range(1, 5):
        style_name = f"Heading {level}"
        try:
            hs = doc.styles[style_name]
            hs.font.name = FONT_HEADING
            hs.font.size = HEADING_SIZES.get(level, Pt(12))
            hs.font.color.rgb = HEADING_COLORS.get(level, RGBColor(0, 0, 0))
            hs.font.bold = True
            hs.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
            hs.paragraph_format.space_after = Pt(8)
        except KeyError:
            pass

    try:
        bullet_style = doc.styles["List Bullet"]
        bullet_style.font.name = FONT_BODY
        bullet_style.font.size = BODY_SIZE
        bullet_style.font.color.rgb = BODY_COLOR
        bullet_style.paragraph_format.space_after = Pt(4)
        bullet_style.paragraph_format.line_spacing = 1.35
    except KeyError:
        pass

    try:
        number_style = doc.styles["List Number"]
        number_style.font.name = FONT_BODY
        number_style.font.size = BODY_SIZE
        number_style.font.color.rgb = BODY_COLOR
        number_style.paragraph_format.space_after = Pt(4)
        number_style.paragraph_format.line_spacing = 1.35
    except KeyError:
        pass

    return doc


def _add_inline_runs(paragraph, text: str):
    parts = re.split(r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
            run.font.size = BODY_SIZE
            run.font.name = FONT_BODY
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = BODY_SIZE
            run.font.name = FONT_BODY
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.size = BODY_SIZE
            run.font.name = FONT_BODY
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = FONT_CODE
            run.font.size = Pt(10)
            run.font.color.rgb = CODE_COLOR
        else:
            run = paragraph.add_run(part)
            run.font.size = BODY_SIZE
            run.font.name = FONT_BODY


def _render_markdown_to_doc(markdown_text: str) -> Document:
    doc = _setup_document()

    lines = markdown_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        heading_match = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            title = re.sub(r"\*\*|__", "", title)
            doc.add_heading(title, level=level)
            i += 1
            continue

        if re.match(r"^\s*```", line):
            code_lines = []
            i += 1
            while i < len(lines) and not re.match(r"^\s*```", lines[i]):
                code_lines.append(lines[i])
                i += 1
            i += 1
            if code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run("\n".join(code_lines))
                run.font.name = FONT_CODE
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        unordered_match = re.match(r"^(\s*)([-*+])\s+(.+)$", line)
        if unordered_match:
            indent_level = len(unordered_match.group(1)) // 2
            text = unordered_match.group(3).strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(1.27 * (indent_level + 1))
            _add_inline_runs(p, text)
            i += 1
            continue

        ordered_match = re.match(r"^(\s*)(\d+)\.\s+(.+)$", line)
        if ordered_match:
            indent_level = len(ordered_match.group(1)) // 2
            text = ordered_match.group(3).strip()
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(1.27 * (indent_level + 1))
            _add_inline_runs(p, text)
            i += 1
            continue

        if line.strip() == "":
            i += 1
            continue

        para_lines = [line]
        i += 1
        while i < len(lines):
            next_line = lines[i]
            if (re.match(r"^#{1,4}\s+", next_line)
                    or re.match(r"^\s*[-*+]\s+", next_line)
                    or re.match(r"^\s*\d+\.\s+", next_line)
                    or re.match(r"^\s*```", next_line)
                    or next_line.strip() == ""):
                break
            para_lines.append(next_line)
            i += 1

        full_text = " ".join(l.strip() for l in para_lines)
        p = doc.add_paragraph()
        _add_inline_runs(p, full_text)

    return doc


async def export_markdown_to_word(markdown_text: str) -> str:
    if not markdown_text or not markdown_text.strip():
        logger.warning("📄 [WordExport] Markdown 文本为空，跳过导出")
        return ""

    output_dir = os.path.join("temp_workspace", "output")
    os.makedirs(output_dir, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Deep_Research_Report_{timestamp}.docx"
    output_path = os.path.join(output_dir, filename)

    def _sync_export():
        doc = _render_markdown_to_doc(markdown_text)
        doc.save(output_path)
        return output_path

    try:
        result_path = await asyncio.to_thread(_sync_export)
        logger.info(f"📄 [WordExport] 导出成功: {result_path} ({os.path.getsize(result_path)} bytes)")
        return result_path
    except Exception as e:
        logger.error(f"📄 [WordExport] 导出失败: {e}")
        return ""
