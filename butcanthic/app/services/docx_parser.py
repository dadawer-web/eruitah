"""
DOCX 异步解析器
融合原 docx_table_extractor.py + xml_processor.py + table_structure_fixer.py 核心逻辑
全异步接口，无 PyQt5 痕迹
"""

import asyncio
import base64
import logging
import os
import re
import subprocess
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional
from xml.etree import ElementTree as ET

from bs4 import BeautifulSoup
from lxml import etree

logger = logging.getLogger(__name__)

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS_MAP = {"w": _W_NS}

_MAMMOTH_SCRIPT = str(Path(__file__).resolve().parent / "docx_to_html_converter.js")


class TableStructureFixer:
    """HTML 表格结构修复器"""

    def fix_table_structure(self, html_content: str) -> str:
        try:
            soup = BeautifulSoup(html_content, "html.parser")
            table = soup.find("table")
            if not table:
                return html_content

            structure = self._analyze(table)
            fixed = self._fix(structure)
            return self._rebuild(fixed, table.attrs)
        except Exception as e:
            logger.warning(f"Table fix failed: {e}")
            return html_content

    def _analyze(self, table) -> Dict:
        rows = table.find_all("tr")
        structure = {"rows": [], "max_logical_cols": 0}
        for row_idx, row in enumerate(rows):
            cells = row.find_all(["td", "th"])
            row_info = {"row_idx": row_idx, "cells": [], "logical_width": 0}
            logical_col = 0
            for cell in cells:
                colspan = int(cell.get("colspan", 1))
                rowspan = int(cell.get("rowspan", 1))
                text = cell.get_text(strip=True)
                row_info["cells"].append(
                    {
                        "logical_col": logical_col,
                        "colspan": colspan,
                        "rowspan": rowspan,
                        "text": text,
                        "is_empty": not text,
                    }
                )
                logical_col += colspan
            row_info["logical_width"] = logical_col
            structure["rows"].append(row_info)
            structure["max_logical_cols"] = max(structure["max_logical_cols"], logical_col)
        return structure

    def _fix(self, structure: Dict) -> Dict:
        grid = self._build_grid(structure)
        for r_idx, row in enumerate(grid):
            for c_idx, cell in enumerate(row):
                if cell is None:
                    grid[r_idx][c_idx] = {
                        "text": "",
                        "is_main": True,
                        "colspan": 1,
                        "rowspan": 1,
                    }
        return {
            "grid": grid,
            "rows": len(grid),
            "cols": len(grid[0]) if grid else 0,
        }

    def _build_grid(self, structure: Dict) -> List[List]:
        rows_count = len(structure["rows"])
        cols_count = structure["max_logical_cols"]
        grid = [[None for _ in range(cols_count)] for _ in range(rows_count)]
        for row_info in structure["rows"]:
            r = row_info["row_idx"]
            logical_col = 0
            for cell_info in row_info["cells"]:
                while logical_col < cols_count and grid[r][logical_col] is not None:
                    logical_col += 1
                if logical_col >= cols_count:
                    break
                cs = cell_info["colspan"]
                rs = cell_info["rowspan"]
                for dr in range(min(rs, rows_count - r)):
                    for dc in range(min(cs, cols_count - logical_col)):
                        grid[r + dr][logical_col + dc] = {
                            "text": cell_info["text"] if dr == 0 and dc == 0 else "",
                            "is_main": dr == 0 and dc == 0,
                            "colspan": cs if dr == 0 and dc == 0 else 1,
                            "rowspan": rs if dr == 0 and dc == 0 else 1,
                        }
                logical_col += cs
        return grid

    def _rebuild(self, structure: Dict, table_attrs: Dict) -> str:
        grid = structure["grid"]
        attrs_str = "".join(f' {k}="{v}"' for k, v in table_attrs.items())
        parts = [f"<table{attrs_str}>"]
        for row in grid:
            parts.append("<tr>")
            for cell in row:
                if cell and cell.get("is_main"):
                    ca = []
                    if cell.get("colspan", 1) > 1:
                        ca.append(f'colspan="{cell["colspan"]}"')
                    if cell.get("rowspan", 1) > 1:
                        ca.append(f'rowspan="{cell["rowspan"]}"')
                    attr = " " + " ".join(ca) if ca else ""
                    text = cell.get("text", "").replace("\n", "<br>")
                    parts.append(f"<td{attr}>{text}</td>")
            parts.append("</tr>")
        parts.append("</table>")
        return "".join(parts)


class DocxParser:
    """DOCX 异步解析器 - 表格提取 + XML 处理"""

    def __init__(self):
        self._fixer = TableStructureFixer()

    async def extract_tables(self, file_path: str) -> List[Dict[str, Any]]:
        """
        从 DOCX 中提取所有表格 (XML + HTML)

        Returns:
            [{"index": int, "xml_content": str, "html_content": str,
              "row_count": int, "col_count": int}]
        """
        xml_tables, html_tables = await asyncio.gather(
            asyncio.to_thread(self._extract_xml_tables, file_path),
            asyncio.to_thread(self._extract_html_tables, file_path),
        )
        return self._merge(xml_tables, html_tables)

    async def convert_docx_to_html(self, file_path: str) -> str:
        """使用 mammoth.js 将整个 DOCX 转为 HTML"""
        return await asyncio.to_thread(self._run_mammoth, file_path)

    async def extract_images(self, file_path: str) -> List[str]:
        """
        从 DOCX 中提取所有内嵌图片，返回 Base64 编码的 data URI 列表

        Returns:
            ["data:image/png;base64,...", "data:image/jpeg;base64,...", ...]
        """
        return await asyncio.to_thread(self._extract_images_sync, file_path)

    async def extract_text_with_images(self, file_path: str) -> Dict[str, Any]:
        """
        提取 DOCX 的文本和图片，在文本中插入 [Image_ID: N] 占位符

        Returns:
            {"text": "带占位符的文本", "images": ["data:image/png;base64,...", ...]}
        """
        return await asyncio.to_thread(self._extract_text_with_images_sync, file_path)

    def _extract_text_with_images_sync(self, file_path: str) -> Dict[str, Any]:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        images = []
        text_parts = []
        image_idx = 0

        try:
            doc = Document(file_path)

            rel_id_to_idx = {}
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    rel_id_to_idx[rel.rId] = image_idx
                    try:
                        blob = rel.target_part.blob
                        ct = rel.target_part.content_type
                        if not ct:
                            ct = "image/png"
                        b64 = base64.b64encode(blob).decode("utf-8")
                        data_uri = f"data:{ct};base64,{b64}"
                        images.append(data_uri)
                        image_idx += 1
                    except Exception as e:
                        logger.warning(f"DOCX: failed to extract image rel {rel.rId}: {e}")

            for para in doc.paragraphs:
                para_text = ""
                for run in para.runs:
                    drawing_elems = run._element.findall(
                        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                    )
                    inline_elems = run._element.findall(
                        ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
                    )

                    found_image = False
                    for drawing in drawing_elems:
                        blip = drawing.find(
                            ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                        )
                        if blip is not None:
                            embed = blip.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                            )
                            if embed and embed in rel_id_to_idx:
                                idx = rel_id_to_idx[embed]
                                para_text += f" [Image_ID: {idx}] "
                                found_image = True

                    for inline in inline_elems:
                        blip = inline.find(
                            ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                        )
                        if blip is not None:
                            embed = blip.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                            )
                            if embed and embed in rel_id_to_idx:
                                idx = rel_id_to_idx[embed]
                                para_text += f" [Image_ID: {idx}] "
                                found_image = True

                    if run.text:
                        para_text += run.text

                if para_text.strip():
                    text_parts.append(para_text.strip())

            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        text_parts.append(" | ".join(row_texts))

        except Exception as e:
            logger.warning(f"DOCX text+image extraction failed: {e}")

        return {"text": "\n".join(text_parts), "images": images}

    def _extract_images_sync(self, file_path: str) -> List[str]:
        images = []
        try:
            with zipfile.ZipFile(file_path, "r") as zf:
                for name in zf.namelist():
                    if name.startswith("word/media/"):
                        data = zf.read(name)
                        ext = Path(name).suffix.lower()
                        mime_map = {
                            ".png": "image/png",
                            ".jpg": "image/jpeg",
                            ".jpeg": "image/jpeg",
                            ".gif": "image/gif",
                            ".bmp": "image/bmp",
                            ".tiff": "image/tiff",
                            ".tif": "image/tiff",
                            ".emf": "image/x-emf",
                            ".wmf": "image/x-wmf",
                        }
                        mime = mime_map.get(ext, "image/png")
                        if mime in ("image/x-emf", "image/x-wmf"):
                            continue
                        b64 = base64.b64encode(data).decode("utf-8")
                        data_uri = f"data:{mime};base64,{b64}"
                        images.append(data_uri)
        except Exception as e:
            logger.warning(f"DOCX image extraction failed: {e}")
        return images

    async def replace_tables_in_docx(
        self, docx_path: str, replacements: List[Dict[str, Any]], output_path: str
    ) -> str:
        """
        替换 DOCX 中的表格 XML 并保存到新文件

        Args:
            docx_path: 原始 DOCX 路径
            replacements: [{"index": int, "new_xml": str}]
            output_path: 输出路径

        Returns:
            输出文件路径
        """
        import shutil

        await asyncio.to_thread(shutil.copy2, docx_path, output_path)
        await asyncio.to_thread(
            self._do_replace_in_docx, output_path, replacements
        )
        return output_path

    def _extract_xml_tables(self, file_path: str) -> List[Dict]:
        results = []
        try:
            with zipfile.ZipFile(file_path, "r") as zf:
                doc_xml = zf.read("word/document.xml").decode("utf-8")
            root = ET.fromstring(doc_xml)
            tables = root.findall(".//w:tbl", _NS_MAP)
            for i, tbl in enumerate(tables):
                results.append(
                    {
                        "index": i + 1,
                        "xml_content": ET.tostring(tbl, encoding="unicode"),
                        "row_count": len(tbl.findall(".//w:tr", _NS_MAP)),
                        "col_count": self._count_xml_cols(tbl),
                    }
                )
        except Exception as e:
            logger.warning(f"XML table extraction failed: {e}")
        return results

    def _count_xml_cols(self, tbl_element) -> int:
        try:
            first_row = tbl_element.find(".//w:tr", _NS_MAP)
            if first_row is None:
                return 0
            total = 0
            for tc in first_row.findall(".//w:tc", _NS_MAP):
                gs = tc.find(".//w:gridSpan", _NS_MAP)
                if gs is not None:
                    val = gs.get(f"{{{_W_NS}}}val", "1")
                    total += int(val)
                else:
                    total += 1
            return total
        except Exception:
            return 0

    def _extract_html_tables(self, file_path: str) -> List[Dict]:
        results = []
        try:
            html = self._run_mammoth(file_path)
            if not html:
                return results
            for i, m in enumerate(
                re.finditer(r"<table[^>]*>.*?</table>", html, re.DOTALL | re.IGNORECASE)
            ):
                raw = m.group(0)
                fixed = self._fixer.fix_table_structure(raw)
                results.append(
                    {
                        "index": i + 1,
                        "html_content": fixed,
                        "row_count": self._count_html_rows(fixed),
                        "col_count": self._count_html_cols(fixed),
                    }
                )
        except Exception as e:
            logger.warning(f"HTML table extraction failed: {e}")
        return results

    def _run_mammoth(self, file_path: str) -> str:
        script = _MAMMOTH_SCRIPT
        if not os.path.exists(script):
            raise FileNotFoundError(
                f"mammoth.js converter not found at: {script}. "
                f"Please run 'npm install' in {os.path.dirname(script)}"
            )

        abs_file_path = str(Path(file_path).resolve())
        if not os.path.exists(abs_file_path):
            raise FileNotFoundError(f"DOCX file not found: {abs_file_path}")

        logger.info(f"_run_mammoth: script={script}, input={abs_file_path}")

        try:
            result = subprocess.run(
                ["node", script, abs_file_path],
                capture_output=True,
                text=True,
                encoding="utf-8",
                timeout=60,
                cwd=str(Path(script).parent),
            )
        except subprocess.TimeoutExpired:
            raise RuntimeError("mammoth.js conversion timed out (60s)")
        except FileNotFoundError:
            raise RuntimeError("Node.js not found. Please install Node.js and add it to PATH")

        if result.returncode != 0:
            stderr = result.stderr.strip() if result.stderr else "unknown error"
            logger.error(f"_run_mammoth: node process failed (rc={result.returncode}), stderr:\n{stderr}")
            raise RuntimeError(f"mammoth.js failed: {stderr}")

        if result.stderr:
            stderr_lines = result.stderr.strip()
            if "警告" in stderr_lines or "warning" in stderr_lines.lower():
                logger.info(f"_run_mammoth: conversion warnings:\n{stderr_lines}")

        html = result.stdout.strip()
        logger.info(f"NodeJS 执行完毕，获取到的 HTML 长度为: {len(html)}")

        if not html:
            logger.error(
                f"_run_mammoth: empty HTML output for {os.path.basename(file_path)}, "
                f"stderr was: {result.stderr.strip() if result.stderr else '(empty)'}"
            )
            raise ValueError(
                f"mammoth.js returned empty HTML for {os.path.basename(file_path)}. "
                "The DOCX may be empty or contain no convertible content."
            )

        return html

    @staticmethod
    def _count_html_rows(html: str) -> int:
        return len(re.findall(r"<tr[^>]*>", html, re.IGNORECASE))

    @staticmethod
    def _count_html_cols(html: str) -> int:
        m = re.search(r"<tr[^>]*>(.*?)</tr>", html, re.DOTALL | re.IGNORECASE)
        if not m:
            return 0
        total = 0
        for cm in re.finditer(
            r"<(td|th)([^>]*?)>", m.group(1), re.IGNORECASE
        ):
            cs = re.search(r'colspan\s*=\s*["\']?(\d+)', cm.group(2), re.IGNORECASE)
            total += int(cs.group(1)) if cs else 1
        return total

    @staticmethod
    def _merge(
        xml_tables: List[Dict], html_tables: List[Dict]
    ) -> List[Dict[str, Any]]:
        merged = []
        for i in range(max(len(xml_tables), len(html_tables))):
            xd = xml_tables[i] if i < len(xml_tables) else {}
            hd = html_tables[i] if i < len(html_tables) else {}
            merged.append(
                {
                    "index": hd.get("index", xd.get("index", i + 1)),
                    "xml_content": xd.get("xml_content", ""),
                    "html_content": hd.get("html_content", ""),
                    "row_count": hd.get("row_count", xd.get("row_count", 0)),
                    "col_count": hd.get("col_count", xd.get("col_count", 0)),
                }
            )
        return merged

    def _do_replace_in_docx(
        self, docx_path: str, replacements: List[Dict[str, Any]]
    ):
        with zipfile.ZipFile(docx_path, "r") as zf:
            doc_xml = zf.read("word/document.xml").decode("utf-8")

        parser = etree.XMLParser(recover=True)
        root = etree.fromstring(doc_xml.encode("utf-8"), parser)
        existing = root.xpath(".//w:tbl", namespaces=_NS_MAP)

        for rep in replacements:
            idx = rep.get("index", 1) - 1
            new_xml = rep.get("new_xml", "")
            if not new_xml or idx >= len(existing):
                continue
            try:
                new_tbl = etree.fromstring(new_xml.encode("utf-8"), parser)
                parent = existing[idx].getparent()
                if parent is not None:
                    parent.replace(existing[idx], new_tbl)
            except Exception as e:
                logger.warning(f"Table {idx + 1} replacement failed: {e}")

        new_doc = etree.tostring(root, xml_declaration=True, encoding="UTF-8")
        with zipfile.ZipFile(docx_path, "a") as zf:
            zf.writestr("word/document.xml", new_doc.decode("utf-8"))


class XmlProcessor:
    """XML 处理器 - 命名空间统一化、清理、验证、分块"""

    def normalize_namespace(self, xml_str: str) -> str:
        if not xml_str:
            return xml_str
        for i in range(10):
            ns = f"ns{i}"
            xml_str = xml_str.replace(f"<{ns}:", "<w:")
            xml_str = xml_str.replace(f"</{ns}:", "</w:")
            xml_str = re.sub(rf"\s{ns}:", " w:", xml_str)
            xml_str = xml_str.replace(f"xmlns:{ns}=", "xmlns:w=")
        if "xmlns:w=" not in xml_str and "<w:" in xml_str:
            xml_str = xml_str.replace(
                "<w:tbl",
                '<w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
                1,
            )
        return xml_str

    def clean_and_validate(self, xml_content: str) -> Optional[str]:
        if not xml_content:
            return None
        xml_content = self.normalize_namespace(xml_content.strip())
        xml_content = re.sub(r"\s+", " ", xml_content).replace("> <", "><")
        if not xml_content.endswith("</w:tbl>") and "<w:tbl" in xml_content:
            xml_content += "</w:tbl>"
        try:
            etree.fromstring(xml_content.encode("utf-8"))
            return xml_content
        except Exception:
            return None

    def extract_xml_from_response(self, response: str) -> Optional[str]:
        content = response.strip()
        for prefix in ("```xml", "```"):
            if content.startswith(prefix):
                content = content[len(prefix):]
        if content.endswith("```"):
            content = content[:-3]
        content = content.strip()
        start = -1
        for pat in (r"<w:tbl[^>]*>", r"<ns\d*:tbl[^>]*>"):
            m = re.search(pat, content)
            if m:
                start = m.start()
                break
        if start < 0:
            return None
        end = -1
        for pat in (r"</w:tbl>", r"</ns\d*:tbl>"):
            matches = list(re.finditer(pat, content))
            if matches:
                end = matches[-1].end()
                break
        if end > start:
            return content[start:end]
        extracted = content[start:]
        if "</w:tbl>" not in extracted:
            extracted += "</w:tbl>"
        return extracted

    def split_table_chunks(self, table_xml: str, chunk_size: int = 15000) -> List[str]:
        if len(table_xml) <= chunk_size:
            return [table_xml]
        header_m = re.search(r"(<w:tbl[^>]*>.*?</w:tblPr>)", table_xml, re.DOTALL)
        if not header_m:
            header_m = re.search(r"(<w:tbl[^>]*>)", table_xml)
        header = header_m.group(1) if header_m else '<w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        if "xmlns:w=" not in header:
            header = header.replace("<w:tbl", '<w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"', 1)
        rows = re.findall(r"<w:tr[^>]*>.*?</w:tr>", table_xml, re.DOTALL)
        if not rows:
            return [table_xml]
        chunks = []
        current_rows = []
        current_size = len(header) + len("</w:tbl>")
        for row in rows:
            if current_size + len(row) > chunk_size and current_rows:
                chunks.append(header + "".join(current_rows) + "</w:tbl>")
                current_rows = []
                current_size = len(header) + len("</w:tbl>")
            current_rows.append(row)
            current_size += len(row)
        if current_rows:
            chunks.append(header + "".join(current_rows) + "</w:tbl>")
        return chunks

    def merge_table_chunks(self, chunk_results: List[str]) -> str:
        if not chunk_results:
            return ""
        if len(chunk_results) == 1:
            return chunk_results[0]
        header_m = re.search(r"(<w:tbl[^>]*>.*?</w:tblPr>)", chunk_results[0], re.DOTALL)
        if not header_m:
            header_m = re.search(r"(<w:tbl[^>]*>)", chunk_results[0])
        header = header_m.group(1) if header_m else '<w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        if "xmlns:w=" not in header:
            header = header.replace("<w:tbl", '<w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"', 1)
        all_rows = []
        for chunk in chunk_results:
            all_rows.extend(re.findall(r"<w:tr[^>]*>.*?</w:tr>", chunk, re.DOTALL))
        return header + "".join(all_rows) + "</w:tbl>"
